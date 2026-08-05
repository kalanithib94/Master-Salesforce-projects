"""Generate GPTfy_Agent_Org_Deployment_Guide.docx"""
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

for section in doc.sections:
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)


def set_run_font(run, size=11, bold=False):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold


def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    size = 16 if level == 1 else 13 if level == 2 else 12
    for run in h.runs:
        set_run_font(run, size=size, bold=True)
    return h


def add_para(text, bold=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_bullet(text):
    p = doc.add_paragraph(text, style="List Bullet")
    for run in p.runs:
        set_run_font(run, size=11)
    p.paragraph_format.space_after = Pt(3)
    return p


def add_numbered(text):
    p = doc.add_paragraph(text, style="List Number")
    for run in p.runs:
        set_run_font(run, size=11)
    p.paragraph_format.space_after = Pt(3)
    return p


def add_code(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    run.font.size = Pt(9)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F2F2F2")
    shd.set(qn("w:val"), "clear")
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p._p.get_or_add_pPr().append(shd)
    return p


# Title
title = doc.add_heading("GPTfy Agent — Org Deployment Guide", level=0)
for run in title.runs:
    set_run_font(run, size=22, bold=True)

add_para(
    "This guide explains how to deploy the client-package "
    "(Apex handlers + 25 CRM skills + system prompt) into a Salesforce org."
)
add_para(
    "Package location: Deliverables / client-package "
    "(Apex source lives in Deliverables / force-app)",
    size=10,
)

# 1. Overview
add_heading_styled("1. What gets deployed", 1)
add_para("Deployment has two parts:")
add_bullet(
    "Metadata (via package.xml): 10 Apex classes that implement the agent skills."
)
add_bullet(
    "Data (via SeedClientSkills.apex): Agent record, 25 skill prompts, "
    "and agent–skill links."
)
add_para(
    "Skills are data records, not Metadata. They cannot be included in "
    "package.xml alone. You must run the seed script after Apex deploy."
)

add_heading_styled("Package contents", 2)
table = doc.add_table(rows=5, cols=2)
table.style = "Table Grid"
rows = [
    ("File", "Purpose"),
    ("package.xml", "Metadata manifest — lists the 10 Apex classes to deploy"),
    (
        "SeedClientSkills.apex",
        'Creates agent "GPTfy Agent", inserts 25 skills, links them',
    ),
    (
        "GPTfy_Agent_SystemPrompt_v1.3.0_client.txt",
        "System prompt (v1.3.0) trimmed for this skill set",
    ),
    (
        "SyncSystemPromptToGPTfyAgent.ps1",
        "Writes the system prompt onto GPTfy Agent",
    ),
]
for i, (a, b) in enumerate(rows):
    table.rows[i].cells[0].text = a
    table.rows[i].cells[1].text = b
    for cell in table.rows[i].cells:
        for p in cell.paragraphs:
            for run in p.runs:
                set_run_font(run, size=10, bold=(i == 0))

doc.add_paragraph()

# 2. Prerequisites
add_heading_styled("2. Prerequisites", 1)
add_para("Complete all of the following before deploying.")

add_heading_styled("2.1 Org & product", 2)
add_numbered("Target Salesforce org (Sandbox recommended for first install).")
add_numbered(
    "GPTfy managed package (namespace: ccai) installed and licensed in that org."
)
add_numbered(
    "A user with permission to deploy Apex and create GPTfy records "
    "(Agent, Prompt, Agent Skill, Data Extraction Mapping)."
)

add_heading_styled("2.2 Local tools", 2)
add_numbered("Salesforce CLI (sf) installed on your machine.")
add_numbered("Authenticated to the target org, for example:")
add_code("sf org login web --alias <ClientOrg>")
add_numbered(
    "PowerShell available (for the system-prompt sync script on Windows)."
)
add_numbered(
    "Access to the Deliverables folder "
    "(must contain both client-package and force-app)."
)

add_heading_styled("2.3 Data Extraction Mapping (required)", 2)
add_para(
    "Every skill prompt needs a Data Extraction Mapping Id. "
    "Ids are org-specific — never copy a Mapping Id from another org."
)
add_para("Query Mapping Ids in the target org:")
add_code(
    'sf data query --query "SELECT Id, Name FROM '
    'ccai__AI_Data_Extraction_Mapping__c LIMIT 20" --target-org <ClientOrg>'
)
add_para(
    "If no mapping exists, create one in the GPTfy UI first, then re-run "
    "the query and note the Id (starts with a0…)."
)

add_heading_styled("2.4 Checklist before you start", 2)
add_bullet("GPTfy (ccai) installed")
add_bullet("sf CLI logged in to <ClientOrg>")
add_bullet("Data Extraction Mapping Id known")
add_bullet(
    "Working directory = Deliverables "
    "(parent of client-package and force-app)"
)

# 3. Deploy steps
add_heading_styled("3. Deployment steps", 1)
add_para("Run all commands from the Deliverables folder.")

add_heading_styled("Step 1 — Deploy Apex (package.xml)", 2)
add_para("This deploys the 10 Apex classes listed in package.xml:")
add_bullet("AgenticSkillsBase")
add_bullet(
    "AccountAgenticSkillsHandler, ContactAgenticSkillsHandler, "
    "LeadAgenticSkillsHandler"
)
add_bullet("OpportunityAgenticSkillsHandler, CaseAgenticSkillsHandler")
add_bullet(
    "UtilityAgenticSkillsHandler, ActivityAgenticSkillsHandler "
    "(deployed for compile/facade dependency; no Utility/Activity skills are seeded)"
)
add_bullet("GenericAgenticSkillsHandler + GenericAgenticSkillsHandlerTest")
add_para("Command:")
add_code(
    "sf project deploy start --manifest client-package/package.xml "
    "--target-org <ClientOrg>"
)
add_para(
    "Confirm the deploy succeeds (Status: Succeeded) before continuing."
)

add_heading_styled("Step 2 — Seed agent and 25 skills", 2)
add_para(
    "Open client-package/SeedClientSkills.apex and replace the placeholder:"
)
add_code(
    "final String DATA_MAPPING = 'REPLACE_WITH_DATA_EXTRACTION_MAPPING_ID';"
)
add_para("With your org Mapping Id, for example:")
add_code("final String DATA_MAPPING = 'a0xxxxxxxxxxxxxxx';")
add_para(
    "If you leave the REPLACE_ placeholder, the script stops and does not insert data."
)
add_para("Run the seed:")
add_code(
    "sf apex run --file client-package/SeedClientSkills.apex "
    "--target-org <ClientOrg>"
)
add_para("What this does:")
add_bullet(
    'Creates (or reuses) agent named "GPTfy Agent" with Status = Active'
)
add_bullet(
    "Inserts 25 ccai__AI_Prompt__c skill records "
    "(skips names that already exist)"
)
add_bullet("Links skills to the agent via ccai__AI_Agent_Skill__c")
add_para(
    "UI alternative: Developer Console → Debug → Open Execute Anonymous Window "
    "→ paste the file contents → Execute."
)
add_para(
    "Re-running the seed is safe: existing prompt Names are skipped; "
    "missing agent–skill links are added."
)

add_heading_styled("Step 3 — Sync the system prompt", 2)
add_para(
    "Writes GPTfy_Agent_SystemPrompt_v1.3.0_client.txt onto the GPTfy Agent record:"
)
add_code(
    "powershell -ExecutionPolicy Bypass -File "
    "client-package\\SyncSystemPromptToGPTfyAgent.ps1 -TargetOrg <ClientOrg>"
)

add_heading_styled("Step 4 — Verify in GPTfy", 2)
add_bullet('Agent "GPTfy Agent" exists and Status = Active')
add_bullet("25 skills are linked to the agent")
add_bullet('Smoke-test in the agent UI, e.g.: "Find accounts named Acme"')

# 4. Skills list
add_heading_styled("4. Skills included (25)", 1)
add_para("Account (4)", bold=True)
add_bullet(
    "fuzzy_search_accounts, create_account, update_account_fields, "
    "fetch_account_related_lists"
)
add_para("Contact (5)", bold=True)
add_bullet(
    "fuzzy_search_contacts, fetch_contact_details, create_contact, "
    "update_contact_fields, log_contact_activity"
)
add_para("Lead (5)", bold=True)
add_bullet(
    "fuzzy_search_leads, fetch_lead_details, create_lead, "
    "update_lead_fields, log_lead_activity"
)
add_para("Opportunity (6)", bold=True)
add_bullet(
    "fuzzy_search_opportunities, fetch_opportunity_details, create_opportunity, "
    "update_opportunity_fields, log_opportunity_activity, "
    "add_opportunity_line_item"
)
add_para("Case (5)", bold=True)
add_bullet(
    "fuzzy_search_cases, fetch_case_details, create_case, "
    "update_case_fields, close_case"
)

add_heading_styled("Explicitly excluded from this package", 2)
add_bullet("fetch_account_details")
add_bullet("convert_lead")
add_bullet("fetch_opportunity_recent_changes")
add_bullet(
    "All Utility skills (fetch_picklist_values, fetch_user_info, "
    "fetch_session_context, …)"
)
add_bullet(
    "All Activity-only skills (create_task, create_event, "
    "complete_task, fetch_my_open_tasks)"
)

# 5. Important notes
add_heading_styled("5. Important notes", 1)
add_bullet(
    "Mapping Ids and Agent Ids are org-specific. "
    "Do not reuse production Ids in another org."
)
add_bullet(
    "Always run commands from the Deliverables folder "
    "so force-app resolves correctly for deploy."
)
add_bullet(
    "Deploy Apex first, then seed data, then sync the system prompt — "
    "in that order."
)
add_bullet(
    "Maintainers only: _build_seed.py and _build_prompt.py regenerate "
    "seed/prompt files; clients do not need them for install."
)

# 6. Troubleshooting
add_heading_styled("6. Troubleshooting", 1)
add_para("Deploy fails / classes missing", bold=True)
add_bullet(
    "Confirm you are in Deliverables and "
    "force-app/main/default/classes contains the handler .cls files."
)
add_bullet("Confirm API version and org allows Apex deployment.")

add_para("Seed stops immediately", bold=True)
add_bullet(
    "DATA_MAPPING still has REPLACE_… — set a real Mapping Id from this org."
)

add_para("No Mapping records returned", bold=True)
add_bullet("Create a Data Extraction Mapping in GPTfy, then query again.")

add_para("Agent has fewer than 25 skills", bold=True)
add_bullet(
    "Re-run SeedClientSkills.apex after fixing DATA_MAPPING; "
    "missing links are added."
)

add_para("Prompt sync script fails", bold=True)
add_bullet(
    "Confirm sf auth for <ClientOrg> and that agent GPTfy Agent already "
    "exists (run seed first)."
)

doc.add_paragraph()
add_para(
    "Document version: 1.0  |  Matches client-package README "
    "(25 skills, system prompt v1.3.0)",
    size=9,
)

out = (
    r"c:\Users\ADMIN\OneDrive\Desktop\AGENT LIBRARY"
    r"\Deliverables\client-package\GPTfy_Agent_Org_Deployment_Guide.docx"
)
doc.save(out)
print("Wrote:", out)
